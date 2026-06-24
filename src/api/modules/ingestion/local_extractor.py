"""Fast local text extraction — runs in milliseconds for native-text files.

Supports:
  - .txt, .csv  → read directly
  - .docx       → python-docx
  - .xlsx       → openpyxl
  - .pdf        → pymupdf (detects if scanned; returns empty text if so)

Returns (text, is_scanned) where is_scanned=True means CU OCR is needed.
"""
from __future__ import annotations

import csv
import io
import logging
from typing import BinaryIO

logger = logging.getLogger(__name__)

# Threshold: if average chars/page is below this, treat as scanned PDF
_SCANNED_CHARS_PER_PAGE_THRESHOLD = 80


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, bool]:
    """Extract text from a file using fast local methods.

    Returns:
        (text, needs_cu) — text is the extracted string (may be empty),
        needs_cu=True means the file requires CU OCR/analysis.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext == "txt":
            return _extract_txt(file_bytes), False

        if ext == "csv":
            return _extract_csv(file_bytes), False

        if ext == "docx":
            return _extract_docx(file_bytes), False

        if ext == "xlsx":
            return _extract_xlsx(file_bytes), False

        if ext == "pdf":
            return _extract_pdf(file_bytes)

        # Audio/video should skip local extraction and go to CU transcription.
        if ext in ["wav", "mp3", "mp4", "m4a", "aac", "flac", "ogg"]:
            return "", True

        # Images and other formats always need CU
        return "", True

    except Exception as e:
        logger.warning(f"Local extraction failed for {filename}: {e} — will fall back to CU")
        return "", True


# ─────────────────────────────────────────────────────────
# Per-type extractors
# ─────────────────────────────────────────────────────────

def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _extract_csv(file_bytes: bytes) -> str:
    text = _extract_txt(file_bytes)
    try:
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return text
        lines = [", ".join(f"{k}: {v}" for k, v in row.items() if v) for row in rows]
        return "\n".join(lines)
    except Exception:
        return text


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts).strip()
        if text:
            return text
        logger.debug("python-docx extracted empty text; trying OOXML fallback")
    except Exception as e:
        logger.debug(f"python-docx parse failed, trying OOXML fallback: {e}")

    # Fallback: extract plain text from document.xml
    import re
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            if "word/document.xml" not in zf.namelist():
                raise ValueError("Not a valid DOCX: missing word/document.xml")
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        text = re.sub(r"<[^>]+>", " ", xml)
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            return text
        logger.warning(f"OOXML fallback for DOCX returned empty text (file may be corrupted or have no content)")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction failed (both python-docx and OOXML fallback): {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def _extract_xlsx(file_bytes: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(v) for v in row if v is not None)
            if row_text.strip():
                parts.append(row_text)
    wb.close()
    return "\n".join(parts)


def _extract_pdf(file_bytes: bytes) -> tuple[str, bool]:
    """Extract text from PDF using pymupdf. Returns (text, needs_cu)."""
    import fitz  # pymupdf

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page_count = len(doc)
    all_text = []

    for page in doc:
        text = page.get_text("text")
        all_text.append(text)

    doc.close()
    full_text = "\n".join(all_text).strip()

    if page_count == 0:
        return "", True

    avg_chars_per_page = len(full_text) / page_count
    if avg_chars_per_page < _SCANNED_CHARS_PER_PAGE_THRESHOLD:
        return "", True

    return full_text, False
